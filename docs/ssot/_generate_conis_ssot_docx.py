#!/usr/bin/env python3
"""Generate CONIS SSOT DOCX documents (PT-SSOT-01)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm


OUT_DIR = Path(__file__).resolve().parents[1] / "ssot"


def set_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)


def title_page(doc: Document, title: str, subtitle: str, meta: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(14)

    doc.add_paragraph()
    for line in meta:
        m = doc.add_paragraph()
        m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        m.add_run(line).font.size = Pt(10)

    doc.add_page_break()


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            t.rows[r_idx + 1].cells[c_idx].text = value
    doc.add_paragraph()


def build_brief() -> Document:
    doc = Document()
    set_styles(doc)
    title_page(
        doc,
        "CONIS PROJECT BRIEF",
        "Projektový přehled platformy CONIS",
        [
            "SSOT — Project Brief",
            "Verze: 1.0 · Datum: 2026-08-03",
            "Účel: během 10–15 minut vysvětlit celý projekt",
            "Není marketingový materiál",
        ],
    )

    h1(doc, "1. Co je CONIS")
    p(
        doc,
        "CONIS je produktová platforma pro stavební firmy (typicky katalogové / typové domy), "
        "která pomáhá zákazníkovi porozumět produktu natolik, aby se mohl rozhodnout sám.",
    )
    p(
        doc,
        "CONIS není chatbot, není obchodník a není generický CMS. "
        "Jádrem je Decision Experience dodávaná přes Embed Experience a řízená Shared Runtime.",
    )
    bullets(
        doc,
        [
            "Produkční zákaznická Experience: Embed (Embed.mount).",
            "Obchodní a provozní řízení partnerů: Office Studio (interní CONIS).",
            "Partner Studios: Client, Manager, Sales.",
            "Implementační povrch: Builder Studio (interní).",
        ],
    )

    h1(doc, "2. Poslání platformy")
    p(doc, "Rozumět souvislostem. Propojovat člověka, produkt a jejich společný smysl.")
    p(doc, "Pomoci člověku porozumět natolik, aby se mohl rozhodnout sám.")
    p(doc, "Informace ≠ porozumění. CONIS nevytváří šum — vytváří porozumění.")

    h1(doc, "3. Problém, který řeší")
    bullets(
        doc,
        [
            "Zákazník na webu stavební firmy dostává fragmentované informace (galerie, plány, texty), ale chybí mu souvislá rozhodovací cesta.",
            "Obchodní tým ztrácí čas s nevyjasněnými poptávkami a opakovaným vysvětlováním základů.",
            "Firma nemá standardizovaný způsob, jak převést zájem o dům do kvalifikovaného rozhodnutí a follow-up.",
            "Nasazení „dalšího softwaru“ bez provozního modelu partnera obvykle selhává — CONIS proto spojuje Experience, Partner Environment a obchodní provoz.",
        ],
    )

    h1(doc, "4. Hodnota pro stavební firmy")
    bullets(
        doc,
        [
            "Kvalifikovanější poptávky: zákazník prochází Decision Experience (Priority → FAQ → Chat → Audit → Report).",
            "Opakovatelný provoz: Partner Environment se studii Client / Manager / Sales pod brandem partnera.",
            "Obchodní kontrola: Office řídí nabídku, doručení pilotu, follow-up, lifecycle a administraci partnera.",
            "Stabilní jádro: obchodní logika Experience žije ve Shared Runtime — nepřepisuje se v každém partnerovi zvlášť.",
            "Pilotní vstup: nízký práh (balíčky Pilot / Starter / Studio Partner) bez nutnosti stavět vlastní produkt od nuly.",
        ],
    )

    h1(doc, "5. Přehled Studií")
    h2(doc, "5.1 Client Studio")
    p(
        doc,
        "Partner-facing prostředí pro práci se zákaznickou Experience a souvisejícím projektem. "
        "Není referenční produkční aplikací — produkční Experience je Embed. "
        "Client Studio slouží partnerovi / internímu vývoji k provozu a kontrole Experience povrchu.",
    )

    h2(doc, "5.2 Manager Studio")
    p(
        doc,
        "Partner-facing provozní studio pro manažerský pohled na projekty a aktivitu. "
        "Součást Partner Environment; není interní CONIS Office.",
    )

    h2(doc, "5.3 Sales Studio")
    p(
        doc,
        "Partner-facing obchodní studio pro obchodní práci partnera nad Experience a leady. "
        "Součást Partner Environment.",
    )

    h2(doc, "5.4 Office (interní)")
    p(
        doc,
        "Interní obchodní řídicí vrstva CONIS (Office Studio v1.0). "
        "Partner Registry, nabídky a checkout, dokumenty, doručení pilotu, follow-up, "
        "Partner Environment Lifecycle a Partner Administration. "
        "Office není partner-primary vstup a neobsahuje Decision Runtime logiku.",
    )

    h2(doc, "5.5 Builder (interní)")
    p(
        doc,
        "Interní implementační / authoring povrch. Partner do něj nevstupuje jako do primární Experience. "
        "Office může předat práci přes Builder Handoff po obchodní cestě.",
    )

    h1(doc, "6. Jak funguje pilot")
    numbered(
        doc,
        [
            "CONIS v Office založí / spravuje partnera a připraví Partner Environment (branding, sample project, studia, pozvánka).",
            "Partner projde Invitation & NDA, aktivací účtu a Welcome Journey.",
            "Office doručí pilotní balíček (Pilot Delivery) a sleduje Commercial Follow-up.",
            "Po potvrzení nabídky (Pilot Offer & Checkout) partner přechází do dlouhodobého Partner Environment Lifecycle (Active / Suspended / Archived).",
            "Office spravuje partnera administrativně (balíček, licence, kontakt, poznámky) bez mazání dat.",
            "Zákazník partnera prochází Embed Decision Experience nad připraveným objektem / projektem.",
        ],
    )
    p(
        doc,
        "Pilot slouží k validaci produktu v reálném provozu. Ne každá připomínka znamená změnu — "
        "další rozvoj má vycházet z ověřené zpětné vazby (Pilot Feedback Register) a explicitních produktových rozhodnutí.",
    )

    h1(doc, "7. Pilotní balíčky")
    table(
        doc,
        ["Balíček", "Rozsah", "Cena (CZK)", "Trial", "Poznámka"],
        [
            ["Pilot", "1 dům", "4 970", "90 dní", "Vstupní pilot · Embed Experience"],
            ["Starter", "až 3 domy", "14 970", "90 dní", "Doporučený start · plný brand"],
            ["Studio Partner", "dle nabídky", "dle nabídky", "90 dní", "Rozšířená partnerská spolupráce"],
        ],
    )
    p(
        doc,
        "Checkout v Office je MVP bez platební brány. Obchodní uzavření nabídky aktivuje Partner Environment Lifecycle.",
    )

    h1(doc, "8. Architektura platformy (vysoká úroveň)")
    p(doc, "Závazné oddělení vrstev:")
    bullets(
        doc,
        [
            "Office Studio ≠ Builder Studio ≠ Partner Environment ≠ Embed Runtime",
            "Shared Runtime (packages/core, packages/runtime): Decision Session, Interpretation, Story, Outcome, Priority, Decision Signals.",
            "Embed Experience: UX a tok Experience; konzumuje Runtime, nereplikuje business logiku.",
            "Partner Environment: provozní povrch partnera (studia + brand + projekty + lifecycle).",
            "Office: obchodní orchestrace partnerů; nepřepisuje Runtime / Decision Layer.",
        ],
    )
    p(
        doc,
        "Embed First: jediná produkční Experience aplikace je Embed. "
        "Client Studio (Local) je vývojářský / partnerský provozní povrch, ne produktová reference Experience.",
    )

    h1(doc, "9. Aktuální stav projektu")
    bullets(
        doc,
        [
            "Office Studio v1.0 — schválený implementační baseline (větev feature/cap-p04-founding-partner @ 1f5a873).",
            "Uzavřena řada OF-01…OF-07, CS-01, PE-02…PE-12 (Partner Administration).",
            "Partner Journey v1.0: provisioning → NDA → delivery → follow-up → welcome → lifecycle → administration → studia.",
            "Commercial Journey v1.0: Vítejte → Pilotní program → Dokončit objednávku → Platba → CONIS Studio (partner purchase path).",
            "Mimo v1.0: Welcome Experience (rozpracováno, ID TBD), OF-08 Operations/AI (karanténa).",
            "Governance: Pilot Feedback Register jako vstup pro plánování 1.1+.",
        ],
    )

    h1(doc, "10. Nejbližší roadmapa")
    bullets(
        doc,
        [
            "Product Review finalizace Commercial Journey v1.0 (PT-CJ-05) a první pilotní obchodní nasazení.",
            "Pilotní provoz u reálných partnerů a sběr VALIDATED feedbacku.",
            "Append-only EPICy po v1.0 (např. Welcome Experience s novým ID; OF-08 jen po explicitním reopen).",
            "Plánování Roadmap v1.1+ výhradně z validované zpětné vazby nebo explicitního produktového rozhodnutí.",
        ],
    )

    h1(doc, "11. Co si odnést")
    p(
        doc,
        "CONIS = platforma porozumění a rozhodnutí pro stavební firmy, dodávaná jako Embed Experience "
        "nad Shared Runtime, provozovaná v Partner Environment a řízená interně přes Office.",
    )
    p(doc, "Office Studio v1.0 je hotový obchodní baseline. Další rozvoj jen append-only.")

    return doc


def build_roadmap() -> Document:
    doc = Document()
    set_styles(doc)
    title_page(
        doc,
        "PLATFORM ROADMAP",
        "Referenční stav implementace CONIS",
        [
            "SSOT — Platform Roadmap",
            "Verze: 1.1 · Datum: 2026-08-05",
            "Baseline: feature/cap-p04-founding-partner @ 1f5a873",
            "Commercial Journey v1.0 — partner purchase path",
            "Pouze aktuální stav — bez historických poznámek",
        ],
    )

    h1(doc, "1. Účel")
    p(
        doc,
        "Tento dokument je jediný referenční přehled implementace platformy CONIS "
        "pro orientaci týmu. Odpovídá skutečnému stavu k datu baseline.",
    )

    h1(doc, "2. Milníky projektu")
    table(
        doc,
        ["Milník", "Stav", "Význam"],
        [
            ["Shared Runtime + Embed Experience", "Aktivní", "Produkční Decision Experience"],
            ["Identity & Access", "Hotovo", "Pozvánky, NDA, role, studia"],
            ["Office Studio v1.0", "Hotovo / FROZEN", "Obchodní baseline partnerů"],
            ["Partner Environment v1.0", "Hotovo", "Provisioning + Lifecycle + Administration"],
            ["Commercial Journey v1.0", "Hotovo · Product Review", "Partner purchase: Vítejte → … → CONIS Studio"],
            ["Pilot Feedback → v1.1+", "Připraveno", "Další rozvoj řízen validací"],
        ],
    )

    h1(doc, "3. EPIC roadmap — Office Studio v1.0")
    p(doc, "Všechny položky níže jsou uzavřené a pushnuté.")
    table(
        doc,
        ["EPIC", "Název", "Commit", "Stav"],
        [
            ["OF-01", "Office Shell", "6a1a8b5", "Dokončeno"],
            ["OF-02", "Partner Workspace", "b1a2c94", "Dokončeno"],
            ["OF-03", "Sales Workspace", "5044ca2", "Dokončeno"],
            ["OF-04", "Document Workspace", "ef4857e", "Dokončeno"],
            ["OF-05", "Builder Handoff", "3d339fd", "Dokončeno"],
            ["OF-06", "Office Pilot Journey", "dfa124c", "Dokončeno"],
            ["OF-07", "Identity & Access", "3fb1fe9", "Dokončeno"],
            ["CS-01", "Pilot Partner Provisioning", "b4741a0", "Dokončeno"],
            ["PE-02", "Brand Projection", "b7fe956", "Dokončeno"],
            ["PE-03", "Pilot Workspace", "ea0f6cb", "Dokončeno"],
            ["PE-04", "Invitation & NDA", "dede2fd", "Dokončeno"],
            ["PE-05", "Welcome Journey", "8038d04", "Dokončeno"],
            ["PE-06", "Pilot Delivery", "945ba84", "Dokončeno"],
            ["PE-07", "Pilot Delivery Finalize", "9be34a9", "Dokončeno"],
            ["PE-08", "Commercial Follow-up", "21b2706", "Dokončeno"],
            ["PE-10", "Partner Environment Provisioning", "2ea815f", "Dokončeno"],
            ["PE-09", "Pilot Offer & Checkout", "ae3066d", "Dokončeno"],
            ["PE-11", "Partner Environment Lifecycle", "466f7b5", "Dokončeno"],
            ["PE-12", "Partner Administration", "1f5a873", "Dokončeno"],
        ],
    )
    p(
        doc,
        "Skutečné pořadí uzavření na konci řady: PE-10 → PE-09 → PE-11 → PE-12.",
    )

    h1(doc, "4. Stav jednotlivých oblastí")
    h2(doc, "4.1 Dokončeno")
    bullets(
        doc,
        [
            "Office Studio v1.0 (shell, partner, sales, documents, handoff, pilot journey, identity).",
            "Partner Environment: provisioning, branding, pilot workspace, invitation/NDA, delivery, follow-up.",
            "Pilot Offer & Checkout (Pilot / Starter / Studio Partner).",
            "Partner Environment Lifecycle (Active / Suspended / Archived) + studio access.",
            "Partner Administration (profil, balíček, licence, kontakt, poznámky, audit).",
            "Embed Experience + Shared Runtime jako produkční Decision Experience vrstva.",
            "Commercial Journey v1.0 — partner purchase path (Working Terminal preview).",
        ],
    )

    h2(doc, "4.2 Rozpracováno")
    bullets(
        doc,
        [
            "Welcome Experience — dirty working tree; evoluce Welcome Journey; epic ID zatím nepřiřazeno.",
            "PT-CJ-05 Finalization — Product Review hold před freeze commit.",
        ],
    )

    h2(doc, "4.3 Připraveno / Karanténa")
    bullets(
        doc,
        [
            "OF-08 Operations / AI — karanténa; reopen jen explicitním PT.",
            "Roadmap v1.1+ — plánování až z VALIDATED pilot feedbacku nebo explicitního rozhodnutí.",
            "Pilot Feedback Register — governance vstup pro další verze.",
        ],
    )

    h1(doc, "5. Partner Journey (aktuální)")
    numbered(
        doc,
        [
            "Pilot Partner Provisioning / Partner Environment Provisioning",
            "Invitation & NDA",
            "Pilot Delivery (+ Finalize)",
            "Commercial Follow-up",
            "Welcome Journey",
            "Partner Environment Lifecycle",
            "Partner Administration",
            "Client Studio / Manager Studio / Sales Studio",
        ],
    )

    h1(doc, "5A. Commercial Journey v1.0 (partner purchase)")
    p(
        doc,
        "Partnerova obchodní cesta (produkční preview v Office Working Terminal). "
        "Není totéž jako Office Pilot Journey (OF-06) ani PE Partner Journey spine.",
    )
    numbered(
        doc,
        [
            "Vítejte",
            "Pilotní program",
            "Dokončit objednávku",
            "Platba",
            "CONIS Studio",
        ],
    )
    p(
        doc,
        "Commits: PT-CJ-00…PT-CJ-04 closed through 7c26352. "
        "Inventory: docs/architecture/office/COMMERCIAL-JOURNEY-IMPLEMENTATION-INVENTORY-v1.0.md",
    )

    h1(doc, "6. Architektonická kontrola")
    bullets(
        doc,
        [
            "Office ≠ Builder ≠ Partner Environment ≠ Runtime — PASS",
            "Commercial Journey ≠ Office Pilot Journey (OF-06) ≠ PE Partner Journey — PASS",
            "Office neobsahuje Decision Layer / Runtime business logiku.",
            "Další EPICy pouze append-only; historie a číslování v1.0 jsou immutable.",
        ],
    )

    h1(doc, "7. Další kroky")
    numbered(
        doc,
        [
            "Uzavřít Product Review Commercial Journey v1.0 (PT-CJ-05) a připravit první pilotní obchod.",
            "Provozovat Office Studio v1.0 a Partner Environment u pilotních partnerů.",
            "Sbírat feedback do Pilot Feedback Register (NEW → … → VALIDATED).",
            "Uzavřít Welcome Experience pod novým append-only epic ID.",
            "OF-08 neotevírat bez explicitního PT.",
            "Sestavit Roadmap v1.1 pouze z VALIDATED položek / produktových rozhodnutí.",
        ],
    )

    h1(doc, "8. Baseline")
    p(doc, "OFFICE BASELINE v1.0 — FROZEN.")
    p(doc, "Office Studio v1.0 je schválený kanonický implementační baseline.")
    p(doc, "Referenční dokumentace: docs/architecture/office/OFFICE-ROADMAP-v2.0.md")

    return doc


def build_bible() -> Document:
    doc = Document()
    set_styles(doc)
    title_page(
        doc,
        "PRODUCT BIBLE",
        "Produktový SSOT platformy CONIS",
        [
            "SSOT — Product",
            "Verze: 1.0 · Datum: 2026-08-03",
            "Hlavní referenční dokument produktu",
            "Není marketing · není implementační návod",
        ],
    )

    h1(doc, "1. Produktová filozofie")
    p(doc, "CONIS je platforma porozumění mezi člověkem a produktem.")
    bullets(
        doc,
        [
            "Není chatbot.",
            "Není obchodník.",
            "Není poradce v běžném smyslu „doporučovače“.",
            "Cíl: porozumění souvislostem → důvěra → rozhodnutí člověka.",
        ],
    )
    p(doc, "Rozhodovací rovnice: Souvislosti → Porozumění → Důvěra → Rozhodnutí.")
    p(doc, "Charakter: klidný, přesný, sebevědomý, kultivovaný, trpělivý.")

    h1(doc, "2. Terminologie")
    table(
        doc,
        ["Termín", "Význam"],
        [
            ["CONIS", "Produktová platforma (identita + Experience + provoz)"],
            ["Embed", "Produkční zákaznická Experience (Embed.mount)"],
            ["Shared Runtime", "SSOT business logiky Decision Experience"],
            ["Decision Experience", "Tok Priority → FAQ → Chat → Audit → Report"],
            ["Partner Environment", "Provozní prostředí partnera (studia, brand, projekty, lifecycle)"],
            ["Partner Workspace", "Modul Office (OF-02) — registry partnerů, ne Partner Environment"],
            ["Office Studio", "Interní obchodní řídicí vrstva CONIS"],
            ["Builder Studio", "Interní implementační / authoring povrch"],
            ["Client / Manager / Sales Studio", "Partner-facing studia"],
            ["Pilot", "Validace produktu v reálném partner provozu"],
            ["Office Baseline v1.0", "Zmrazený implementační baseline Office"],
        ],
    )

    h1(doc, "3. Architektura produktu")
    bullets(
        doc,
        [
            "Dvě produkční vrstvy Experience: Shared Runtime + Embed Experience.",
            "Office orchestrace partnerů je oddělená od Runtime.",
            "Partner Environment je partner surface, ne čtvrtý Runtime.",
            "Zákaz: míchat Office, Builder, Partner Environment a Runtime do jedné „aplikace“.",
        ],
    )
    p(doc, "Embed First: Embed je jediná produkční Experience aplikace.")

    h1(doc, "4. Studia")
    h2(doc, "Partner-facing")
    bullets(
        doc,
        [
            "Client Studio — provoz Experience / projektu partnera.",
            "Manager Studio — manažerský provoz.",
            "Sales Studio — obchodní provoz partnera.",
        ],
    )
    h2(doc, "Interní CONIS")
    bullets(
        doc,
        [
            "Office Studio — obchod, delivery, lifecycle, administrace partnerů.",
            "Builder Studio — implementace / handoff; není partner-primary.",
        ],
    )

    h1(doc, "5. Partner Environment")
    p(
        doc,
        "Partner Environment je dlouhodobě spravovaná entita partnera: společnost, logo, hero, "
        "studio branding, licence, projekty, environment status.",
    )
    bullets(
        doc,
        [
            "Provisioning: připravené prostředí + pozvánka.",
            "Lifecycle: Active → Suspended → Archived (Pilot není lifecycle stav).",
            "Studio access Client/Manager/Sales podle lifecycle; Office a Builder zůstávají interní.",
            "Administration: změna balíčku, licence, kontaktu, interní poznámky + audit timeline.",
        ],
    )

    h1(doc, "6. Decision Experience")
    p(
        doc,
        "Decision Experience je zákaznický rozhodovací tok nad objektem (typicky dům). "
        "Business logika patří Runtime. Experience vrstva řeší prezentaci a interakci.",
    )
    bullets(
        doc,
        [
            "Priority — priority a rozhodnutí domácnosti.",
            "FAQ — strukturované otázky.",
            "Chat — dialog v mezích Experience.",
            "Audit — kontrola rozhodnutí.",
            "Report — výstup pro zákazníka / follow-up.",
        ],
    )
    p(
        doc,
        "Detailní produktová gramatika Experience: Decision Experience Grammar (DEG) v docs/product.",
    )

    h1(doc, "7. Obchodní model")
    bullets(
        doc,
        [
            "Partner (stavební firma) provozuje Experience pod vlastním brandem.",
            "CONIS dodává platformu, Runtime, Embed delivery a Office provoz.",
            "Hodnota: kvalitnější rozhodovací cesta a opakovatelný partnerský provoz — ne „prodej softwarové licence jako cíle“.",
            "Office v1.0 pokrývá obchodní uzavření nabídky (checkout MVP bez platební brány).",
        ],
    )

    h1(doc, "8. Pilotní model")
    bullets(
        doc,
        [
            "Balíčky: Pilot (1 dům), Starter (až 3 domy), Studio Partner.",
            "Cíl pilotu: validace produktu, ne nekonečný backlog nápadů.",
            "Zpětná vazba: Pilot Feedback Register → VALIDATED → roadmapa 1.1+.",
            "Ne každá připomínka se implementuje.",
        ],
    )

    h1(doc, "9. Design principy")
    bullets(
        doc,
        [
            "Simplifikace před rychlostí implementace.",
            "Jedna Experience — Embed jako kanonická.",
            "Žádná duplicita Runtime logiky v UI vrstvách.",
            "Kanonické názvy; žádné oživování retired aliasů.",
            "Append-only rozvoj po zmrazeném baseline.",
        ],
    )

    h1(doc, "10. UX principy")
    bullets(
        doc,
        [
            "Klid a přesnost před agresivním prodejem.",
            "Rozhodnutí zůstává u člověka.",
            "Jeden účel sekce / obrazovky.",
            "Partner brand v Partner Environment; CONIS řídí Office interně.",
            "Experience nesmí vypadat jako dashboard admin nástroje.",
        ],
    )

    h1(doc, "11. Omezení systému")
    bullets(
        doc,
        [
            "Office neimplementuje Decision Layer / Runtime.",
            "Builder není partnerský vstup Experience.",
            "Partner Workspace ≠ Partner Environment.",
            "OF-08 Operations/AI je mimo v1.0 (karanténa).",
            "Checkout bez platební brány / fakturace v Office v1.0.",
            "Zmrazené baseline a číslování se nepřepisují.",
        ],
    )

    h1(doc, "12. Budoucí směr")
    bullets(
        doc,
        [
            "Provozovat pilot u reálných partnerů.",
            "Řídit v1.1+ přes validovaný feedback a explicitní produktová rozhodnutí.",
            "Dokončit Welcome Experience jako append-only epic.",
            "Reopen OF-08 jen při jasné potřebě a PT.",
            "Udržet Embed First a oddělení vrstev při každém novém EPIC.",
        ],
    )

    h1(doc, "13. Autorita dokumentu")
    p(
        doc,
        "PRODUCT BIBLE je hlavní SSOT produktu CONIS. "
        "Při konfliktu s marketingovým textem platí tento dokument a CONIS Constitution. "
        "Při konfliktu s implementačním stavem Office platí OFFICE ROADMAP v2.0 / Office Baseline v1.0. "
        "Při konfliktu o Experience gramatice platí DEG.",
    )

    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    brief = build_brief()
    brief_path = OUT_DIR / "CONIS PROJECT BRIEF.docx"
    brief.save(brief_path)

    roadmap = build_roadmap()
    roadmap_path = OUT_DIR / "PLATFORM ROADMAP.docx"
    roadmap.save(roadmap_path)

    bible = build_bible()
    bible_path = OUT_DIR / "PRODUCT BIBLE.docx"
    bible.save(bible_path)

    print(f"Wrote: {brief_path}")
    print(f"Wrote: {roadmap_path}")
    print(f"Wrote: {bible_path}")


if __name__ == "__main__":
    main()
